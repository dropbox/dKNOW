#!/usr/bin/env python3
"""
Generate test .doc files for DOC backend testing.

This script creates 5 diverse Microsoft Word 97-2003 binary format (.doc) files
for testing the DOC to DOCX conversion pipeline.

Requirements:
- python-docx: pip install python-docx
- LibreOffice or MS Word to save as .doc format

Since python-docx only writes DOCX format, this script creates .docx files first,
then provides instructions for converting them to .doc using LibreOffice.

Usage:
    python3 CREATE_DOC_TEST_FILES.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_simple_text_doc():
    """1. Simple text with basic formatting"""
    doc = Document()

    # Title
    title = doc.add_heading('Simple Text Document', 0)

    # Paragraphs
    doc.add_paragraph('This is a simple text document with basic formatting.')
    doc.add_paragraph('It contains multiple paragraphs to test text extraction.')

    # List
    doc.add_heading('Features:', level=2)
    doc.add_paragraph('Basic paragraphs', style='List Bullet')
    doc.add_paragraph('Headings at multiple levels', style='List Bullet')
    doc.add_paragraph('Simple lists', style='List Bullet')

    # More text
    doc.add_paragraph('This document is intentionally simple to verify basic conversion functionality.')

    return doc

def create_formatted_document_doc():
    """2. Rich formatting: bold, italic, colors, fonts"""
    doc = Document()

    # Title with formatting
    title = doc.add_heading('Formatted Document', 0)

    # Paragraph with mixed formatting
    p = doc.add_paragraph()
    p.add_run('This document contains ').font.size = Pt(12)
    bold_run = p.add_run('bold text')
    bold_run.bold = True
    bold_run.font.size = Pt(12)

    p.add_run(', ')
    italic_run = p.add_run('italic text')
    italic_run.italic = True
    italic_run.font.size = Pt(12)

    p.add_run(', and ')
    color_run = p.add_run('colored text')
    color_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    color_run.font.size = Pt(12)
    p.add_run('.')

    # Different font sizes
    doc.add_heading('Font Sizes', level=2)
    for size in [10, 12, 14, 16, 18]:
        p = doc.add_paragraph(f'This is {size}pt text')
        for run in p.runs:
            run.font.size = Pt(size)

    # Alignment
    doc.add_heading('Text Alignment', level=2)
    for align, name in [(WD_PARAGRAPH_ALIGNMENT.LEFT, 'Left'),
                         (WD_PARAGRAPH_ALIGNMENT.CENTER, 'Center'),
                         (WD_PARAGRAPH_ALIGNMENT.RIGHT, 'Right')]:
        p = doc.add_paragraph(f'{name} aligned text')
        p.alignment = align

    return doc

def create_tables_and_columns_doc():
    """3. Document with tables and merged cells"""
    doc = Document()

    doc.add_heading('Document with Tables', 0)

    # Simple table
    doc.add_heading('Simple Table', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Age'
    header_cells[2].text = 'City'

    # Data rows
    data = [
        ('Alice', '30', 'New York'),
        ('Bob', '25', 'Los Angeles'),
    ]
    for i, (name, age, city) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = name
        row_cells[1].text = age
        row_cells[2].text = city

    # Larger table
    doc.add_paragraph()  # Spacing
    doc.add_heading('Larger Table', level=2)
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Medium Shading 1 Accent 1'

    # Header
    header = table2.rows[0].cells
    for i, title in enumerate(['Product', 'Quantity', 'Price', 'Total']):
        header[i].text = title

    # Data
    products = [
        ('Laptop', '2', '$1000', '$2000'),
        ('Mouse', '5', '$20', '$100'),
        ('Keyboard', '3', '$50', '$150'),
        ('Monitor', '1', '$300', '$300'),
    ]
    for i, (product, qty, price, total) in enumerate(products, start=1):
        cells = table2.rows[i].cells
        cells[0].text = product
        cells[1].text = qty
        cells[2].text = price
        cells[3].text = total

    return doc

def create_images_and_objects_doc():
    """4. Document with embedded images placeholder (text description)"""
    doc = Document()

    doc.add_heading('Document with Images and Objects', 0)

    doc.add_paragraph('This document would normally contain embedded images.')
    doc.add_paragraph('[IMAGE: Company Logo]')

    doc.add_heading('Caption Test', level=2)
    doc.add_paragraph('[IMAGE: Product Photo]')
    doc.add_paragraph('Figure 1: Product photo showing the new design', style='Caption')

    doc.add_heading('Shapes and Objects', level=2)
    doc.add_paragraph('[SHAPE: Blue Rectangle]')
    doc.add_paragraph('[SHAPE: Green Circle]')

    return doc

def create_complex_academic_doc():
    """5. Complex academic paper with footnotes, bibliography, figures"""
    doc = Document()

    # Title page
    title = doc.add_heading('The Impact of Document Conversion Technologies', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    author = doc.add_paragraph('by Dr. John Smith')
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    date = doc.add_paragraph('November 2025')
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This paper examines the evolution of document conversion technologies, '
        'focusing on the transition from legacy binary formats to modern XML-based '
        'standards. We analyze the technical challenges involved in parsing and '
        'converting Microsoft Word 97-2003 binary format (.doc) files.'
    )

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Document formats have evolved significantly over the past three decades. '
        'The Microsoft Word 97-2003 binary format, based on the OLE2/CFB '
        '(Compound File Binary) container structure, represented the dominant '
        'word processing format for over a decade.'
    )

    # Methodology
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        'Our research employed a conversion-based approach using platform-native '
        'tools to transform binary .doc files into XML-based .docx format.'
    )

    # Results table
    doc.add_heading('2.1 Conversion Results', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Format'
    header[1].text = 'Success Rate'
    header[2].text = 'Avg. Time (ms)'

    data = [
        ('Simple text', '100%', '250'),
        ('Formatted', '98%', '320'),
        ('With tables', '95%', '450'),
    ]
    for i, (fmt, rate, time) in enumerate(data, start=1):
        cells = table.rows[i].cells
        cells[0].text = fmt
        cells[1].text = rate
        cells[2].text = time

    # Footnotes (simulated as text)
    doc.add_paragraph()
    doc.add_paragraph('¹ Microsoft Corporation, "MS-DOC Specification", 2008')
    doc.add_paragraph('² Apple Inc., "textutil man page", macOS 14.0')

    # Bibliography
    doc.add_page_break()
    doc.add_heading('References', level=1)
    doc.add_paragraph('[1] Microsoft Corporation. MS-DOC: Word 97-2003 Binary File Format Specification. 2008.')
    doc.add_paragraph('[2] ISO/IEC 29500:2016. Office Open XML File Formats. International Organization for Standardization.')
    doc.add_paragraph('[3] ECMA-376. Office Open XML File Formats. ECMA International. 2016.')

    return doc

def main():
    """Create all test .doc files"""

    # Create DOCX files first
    docx_files = {
        'simple_text.docx': create_simple_text_doc(),
        'formatted_document.docx': create_formatted_document_doc(),
        'tables_and_columns.docx': create_tables_and_columns_doc(),
        'images_and_objects.docx': create_images_and_objects_doc(),
        'complex_academic.docx': create_complex_academic_doc(),
    }

    print("Creating DOCX files...")
    for filename, doc in docx_files.items():
        doc.save(filename)
        print(f"  ✓ Created {filename}")

    print("\n" + "="*70)
    print("IMPORTANT: Convert DOCX to DOC format")
    print("="*70)
    print("\nThe .docx files have been created. Now convert them to .doc using LibreOffice:\n")
    print("On macOS/Linux:")
    for filename in docx_files.keys():
        doc_name = filename.replace('.docx', '.doc')
        print(f"  soffice --headless --convert-to doc --outdir . {filename}")

    print("\nAlternatively, use textutil on macOS:")
    for filename in docx_files.keys():
        doc_name = filename.replace('.docx', '.doc')
        print(f"  textutil -convert doc {filename} -output {doc_name}")

    print("\nOr open each .docx file in Microsoft Word and 'Save As' → Word 97-2003 Document (.doc)")
    print("\nAfter conversion, you should have 5 .doc files:")
    for filename in docx_files.keys():
        print(f"  - {filename.replace('.docx', '.doc')}")

    print("\n" + "="*70)

if __name__ == '__main__':
    main()
